#!/usr/bin/env python3
"""Round-trip check a CV the way an applicant tracking system would.

Extracts the text back out of a .pdf or .docx, then reports:
  - whether the text is machine-readable at all (a scanned or image CV is not)
  - whether standard section headings survived
  - whether contact details survived
  - whether every claim you care about survived (--claims)
  - structural risks that break parsers: tables, text boxes, multi-column
    layout, content stranded in headers or footers

Usage:
    python3 ats_check.py CV.pdf
    python3 ats_check.py CV.docx --claims claims.txt
    python3 ats_check.py CV.pdf --claims-inline "EUR 4bn" "14 markets" "Principal"

claims.txt is one claim per line; blank lines and lines starting with # ignored.
Exit code is 0 if everything passed, 1 if anything failed.
"""

import argparse
import os
import re
import subprocess
import sys
import zipfile

SECTIONS = ["experience", "education", "skills"]
SECTION_ALIASES = {
    "experience": ["experience", "employment", "work history", "career"],
    "education": ["education", "qualifications", "academic"],
    "skills": ["skills", "competencies", "expertise", "technical"],
}


def extract_pdf(path):
    try:
        out = subprocess.run(["pdftotext", "-layout", path, "-"],
                             capture_output=True, text=True, timeout=60)
        if out.returncode == 0 and out.stdout.strip():
            return out.stdout
    except (FileNotFoundError, subprocess.SubprocessError):
        pass
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            return "\n".join((p.extract_text() or "") for p in pdf.pages)
    except ImportError:
        pass
    try:
        from pypdf import PdfReader
        return "\n".join((p.extract_text() or "") for p in PdfReader(path).pages)
    except ImportError:
        print("No PDF text extractor available. Install pdfplumber or poppler-utils.")
        sys.exit(2)


def extract_docx(path):
    """Body text only - deliberately ignores headers, footers and text boxes,
    because many parsers do the same."""
    try:
        import docx
        d = docx.Document(path)
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                parts.extend(c.text for c in row.cells)
        return "\n".join(parts)
    except ImportError:
        with zipfile.ZipFile(path) as z:
            xml = z.read("word/document.xml").decode("utf-8", "ignore")
        xml = re.sub(r"</w:p>", "\n", xml)
        return re.sub(r"<[^>]+>", "", xml)


def docx_structure(path):
    risks = []
    try:
        import docx
        d = docx.Document(path)
        if d.tables:
            risks.append(f"{len(d.tables)} table(s) - parsers frequently scramble "
                         "or drop table content. Use plain paragraphs.")
        for section in d.sections:
            head = "\n".join(p.text for p in section.header.paragraphs).strip()
            foot = "\n".join(p.text for p in section.footer.paragraphs).strip()
            if len(head) > 3:
                risks.append(f"Header contains text ({head[:40]!r}...) - many "
                             "parsers drop headers entirely.")
            if len(foot) > 3 and not re.fullmatch(r"[\d\s/of-]+", foot, re.I):
                risks.append(f"Footer contains text ({foot[:40]!r}...) - same risk.")
    except ImportError:
        pass
    try:
        with zipfile.ZipFile(path) as z:
            xml = z.read("word/document.xml").decode("utf-8", "ignore")
        if "<w:cols" in xml and re.search(r'w:num="([2-9])"', xml):
            risks.append("Multi-column layout detected - read order is unreliable. "
                         "Use a single column.")
        if "w:txbxContent" in xml:
            risks.append("Text box detected - content inside text boxes is often "
                         "invisible to parsers.")
    except (zipfile.BadZipFile, KeyError):
        pass
    return risks


def pdf_structure(path, text):
    risks = []
    if len(text.strip()) < 200:
        risks.append("Almost no extractable text - this may be an image or scanned "
                     "PDF. An ATS will see nothing. Export from the source document "
                     "instead.")
    lines = [l for l in text.splitlines() if l.strip()]
    wide_gaps = sum(1 for l in lines if re.search(r"\S {6,}\S", l))
    if lines and wide_gaps / len(lines) > 0.35:
        risks.append("Many lines contain wide internal gaps, which usually means a "
                     "two-column or table layout. Read order may be scrambled.")
    return risks


def norm(s):
    return re.sub(r"[^a-z0-9]+", " ", s.lower()).strip()


def main():
    ap = argparse.ArgumentParser(description="ATS round-trip check for a CV.")
    ap.add_argument("path")
    ap.add_argument("--claims", help="file with one claim per line")
    ap.add_argument("--claims-inline", nargs="*", default=[])
    args = ap.parse_args()

    path = os.path.expanduser(args.path)
    if not os.path.exists(path):
        print(f"Not found: {path}")
        return 2

    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        text = extract_pdf(path)
        risks = pdf_structure(path, text)
    elif ext == ".docx":
        text = extract_docx(path)
        risks = docx_structure(path)
    elif ext in (".txt", ".md"):
        text = open(path, encoding="utf-8").read()
        risks = []
    else:
        print(f"Unsupported file type: {ext}. Use .pdf, .docx, .txt or .md.")
        return 2

    flat = norm(text)
    words = len(text.split())
    failures = []

    print(f"\nATS round-trip: {os.path.basename(path)}")
    print("=" * 60)
    print(f"Extracted {words} words, {len(text.splitlines())} lines.")
    if words < 150:
        failures.append("Very little text extracted - check the file is not an image.")

    print("\nSection headings")
    for key in SECTIONS:
        found = any(norm(a) in flat for a in SECTION_ALIASES[key])
        print(f"  [{'ok' if found else 'MISSING'}] {key}")
        if not found:
            failures.append(f"No recognisable '{key}' heading. Use the standard word.")

    print("\nContact details")
    email = re.search(r"[\w.+-]+@[\w-]+\.[\w.]+", text)
    phone = re.search(r"(\+?\d[\d\s().-]{7,}\d)", text)
    linkedin = "linkedin" in text.lower()
    for label, ok in (("email", bool(email)), ("phone", bool(phone)),
                      ("linkedin", linkedin)):
        print(f"  [{'ok' if ok else 'MISSING'}] {label}")
    if not email:
        failures.append("No email address survived extraction - it may be in a "
                        "header, an image, or a hyperlink with no visible text.")

    print("\nPlaceholders left in the document")
    placeholders = re.findall(r"\[[A-Z][A-Z \d_/-]{3,}\]", text)
    if placeholders:
        for p in sorted(set(placeholders)):
            print(f"  [ACTION] {p}")
        failures.append(f"{len(set(placeholders))} placeholder(s) still in the file.")
    else:
        print("  none")

    claims = list(args.claims_inline)
    if args.claims:
        with open(os.path.expanduser(args.claims), encoding="utf-8") as fh:
            claims += [l.strip() for l in fh
                       if l.strip() and not l.strip().startswith("#")]
    if claims:
        print("\nClaims")
        for c in claims:
            ok = norm(c) in flat
            print(f"  [{'ok' if ok else 'LOST'}] {c}")
            if not ok:
                failures.append(f"Claim did not survive extraction: {c!r}")

    print("\nStructural risks")
    if risks:
        for r in risks:
            print(f"  [WARN] {r}")
    else:
        print("  none detected")

    print("\n" + "=" * 60)
    if failures:
        print(f"FAILED - {len(failures)} issue(s):")
        for f in failures:
            print(f"  - {f}")
        return 1
    print("PASSED - an ATS should read this correctly.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
